from flask import Flask, render_template, Response, request, redirect
from pymysql import connect, cursors, Error
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pandas as pd 
from sqlalchemy import create_engine

app = Flask(__name__)

# Jinja
# blog = {
#     "1": {
#         "id": 1,
#         "title": "Bài viết 1",
#         "content": "ahihi"
#     },
#     "2": {
#         "id": 2,
#         "title": "Bài viết 2",
#         "content": "ahiha"
#     }
# }

config = {
    'host': 'localhost',
    'user': 'root',
    'password': '',
    'database': 'blog',
    }
cnx = connect(**config)

@app.route("/")
def home():
    cur= cnx.cursor()
    sql="Select * from blog_uer"
    cur.execute(sql)
    return render_template("home.html", cur=cur)

       

@app.route('/', methods=["POST"])
def new_blog():
    
    title = request.form["title"]
    content = request.form["content"]
    cur= cnx.cursor()
    sql="INSERT into blog_uer (title,contain,date) VALUES (%s, %s, %s)"
    value=(title,content,datetime.now())
    cur.execute(sql,value)
    cnx.commit()
    return redirect("/", code=302)

@app.route('/blog/<id>',methods=["GET"])
def detail(id):
 
    cur= cnx.cursor()
    sql="select * from blog_uer where id="+str(id)
    cur.execute(sql)
    for i in cur:
        r1 = i[1] 
        r2 = i[2] 
        r0 = i[0] 
    return render_template("update.html",r0=r0, r1=r1,r2=r2)


@app.route("/blog/<id>", methods=["POST"])
def update_blog(id):
    title = request.form["title"]
    content = request.form["content"]
    cur= cnx.cursor()
    sql=f"UPDATE blog_uer SET title='{title}', contain= '{content}' where id= {id}"
    cur.execute(sql)
    cnx.commit()
    return redirect("/", code=302)   


@app.route("/about")
def about():
    return render_template("about.html") 


@app.route("/resignation")
def resignation():
    return render_template("resignation-letter.html")         

@app.route('/resignation', methods=['POST'])    
def sent_resi():
    fullname = request.form["fullname"]
    reason = request.form["reason"]
    document = Document()
    document.add_heading('Document Title', 0)
    document.add_paragraph('A plain paragraph having some ')
    document.save('satic/demo.docx')
    return redirect('/resignation',code=302)

# @app.route('/resignation', methods=['POST'])    
# def sent_resi():
#     fullname = request.form["fullname"]
#     reason = request.form["reason"]
#     file = 'demox.xlsx'
#     writer = pd.ExcelWriter(file, engine='openpyxl')
#     df = pd.DataFrame([
#         [fullname],
#         [reason]
#     ], 
#     columns=['fullname', 'reason'])
#     df.to_excel(writer,  sheet_name='regg', index=False)
#     writer.save()
#     writer.close()
#     return redirect('/resignation',code=302)

# def load_excel():
#     cur= cnx.cursor()
#     df = pd.read_sql('SELECT * FROM blog_uer', cur)
#     df.to_excel('user.xlsx', sheet_name='products', index=False)

if __name__ == "__main__":
    app.run(debug=True)

